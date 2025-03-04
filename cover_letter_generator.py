import PyPDF2
from docx import Document
from nlp_processing import extract_skills_from_description

def generate_cover_letter(job_title, company, job_desc, cv_file):
    # Extract skills from job description
    skills = extract_skills_from_description(job_desc)
    
    # Extract relevant experience from CV
    experience = extract_experience_from_cv(cv_file)

    # Extract name and contact info from the CV
    name, contact_info = extract_name_and_contact_from_cv(cv_file)

    # Create the cover letter template
    cover_letter = f"""
    Dear Hiring Manager,

    I am excited to apply for the {job_title} position at {company}. With a strong background in {', '.join(skills)}, 
    I am eager to contribute my expertise to your team.

    Currently, I am a {experience}. I have successfully contributed to optimizing experimental workflows, improving predictive accuracy, and co-authoring research papers. Furthermore, I have disseminated AI/ML research publications to a community of over 200 data scientists, helping foster collaboration and drive advancements in the field.

    I am excited about the opportunity to apply my technical skills, analytical expertise, and academic background to the {job_title} role at {company}, where I can leverage my experience to help drive positive change in the advertising industry. The opportunity to work on flexible hours and contribute to a team that ensures responsible advertising resonates with my professional values and long-term career goals.

    I look forward to the opportunity to further discuss how my background and skills can contribute to the continued success of {company}. Thank you for your consideration.

    Sincerely,  
    {name}  
    {contact_info}
    """
    
    return cover_letter

# Helper functions for extracting information from CV (if necessary)
def extract_experience_from_cv(cv_file):
    """
    Extracts relevant experience and skills from the uploaded CV.
    This function works for both PDF and DOCX files.
    """
    experience = ""

    # If the CV is a PDF
    if cv_file.name.lower().endswith('.pdf'):
        reader = PyPDF2.PdfReader(cv_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        experience = extract_experience_from_text(text)

    # If the CV is a DOCX
    elif cv_file.name.lower().endswith('.docx'):
        doc = docx.Document(cv_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        experience = extract_experience_from_text(text)

    return experience


def extract_experience_from_text(text):
    """
    Processes the raw text from the CV to extract experience and responsibilities.
    """
    # Simple processing to find sections related to experience
    experience_section = ""
    experience_keywords = ['experience', 'work', 'role', 'responsibilities']

    for line in text.split('\n'):
        for keyword in experience_keywords:
            if keyword in line.lower():
                experience_section += line.strip() + "\n"
                break

    return experience_section


def extract_name_and_contact_from_cv(cv_file):
    """
    Extracts name and contact information from the CV (works for both PDF and DOCX files).
    This assumes the name is at the top and contact information follows.
    """
    name = ""
    contact_info = ""

    # If the CV is a PDF
    if cv_file.name.lower().endswith('.pdf'):
        reader = PyPDF2.PdfReader(cv_file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        name, contact_info = extract_name_and_contact_from_text(text)

    # If the CV is a DOCX
    elif cv_file.name.lower().endswith('.docx'):
        doc = docx.Document(cv_file)
        text = "\n".join([para.text for para in doc.paragraphs])
        name, contact_info = extract_name_and_contact_from_text(text)

    return name, contact_info


def extract_name_and_contact_from_text(text):
    """
    Processes the raw text from the CV to extract name and contact information.
    Assumes name is the first line and contact info follows.
    """
    lines = text.split("\n")
    
    # Assuming the first line is the name and second line contains contact information
    name = lines[0] if len(lines) > 0 else "Your Full Name"
    contact_info = lines[1] if len(lines) > 1 else "Your Contact Information"
    
    return name, contact_info

